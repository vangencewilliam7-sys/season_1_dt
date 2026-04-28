"""
runtime/readers/filesystem_reader.py

Concrete DocumentReader — reads from a local directory of files.

Supports: .txt, .md, .pdf (text extraction), .json, .docx

Use for:
- Development and testing
- KBs that export documents as files
- Master Cases stored as flat files in a folder structure

Folder convention expected:
    <base_dir>/
        <expert_id>/
            knowledge_hub/      ← documents from the Knowledge Hub
                doc1.txt
                doc2.pdf
            master_cases/       ← Master Case documents
                case1.md
                case2.txt
"""

import logging
import os
from pathlib import Path
from typing import List

from runtime.readers.base_reader import BaseDocumentReader, Document

logger = logging.getLogger(__name__)

# Supported file extensions and their reader strategies
SUPPORTED_EXTENSIONS = {".txt", ".md", ".json", ".pdf", ".docx"}


class FilesystemReader(BaseDocumentReader):
    """
    Reads all expert documents from a local directory structure.
    Supports plain text, markdown, JSON, PDF (text), and DOCX files.
    """

    def __init__(self, base_dir: str):
        """
        Args:
            base_dir: Root directory containing per-expert subdirectories.
                      Example: "d:/Expert-Persona/data"
        """
        self.base_dir = Path(base_dir)

    def load(self, expert_id: str) -> List[Document]:
        """
        Load all documents for the given expert from the filesystem.
        Reads both knowledge_hub/ and master_cases/ subdirectories.
        """
        expert_dir = self.base_dir / expert_id

        if not expert_dir.exists():
            logger.warning(f"[FilesystemReader] No directory found for expert: {expert_id} "
                           f"at {expert_dir}")
            return []

        documents = []

        # Read from both subdirectories if present
        for subdir_name in ["knowledge_hub", "master_cases", "."]:
            subdir = expert_dir / subdir_name if subdir_name != "." else expert_dir
            if not subdir.is_dir():
                continue

            for file_path in subdir.rglob("*"):
                if not file_path.is_file():
                    continue
                if file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
                    continue
                # Skip files directly in expert_dir if we already scan subdirs
                if subdir_name == "." and file_path.parent != expert_dir:
                    continue

                doc = self._read_file(file_path, expert_id, subdir_name)
                if doc:
                    documents.append(doc)

        logger.info(f"[FilesystemReader] Loaded {len(documents)} documents "
                    f"for expert: {expert_id}")
        return documents

    def _read_file(self, path: Path, expert_id: str, source_type: str) -> Document | None:
        """Read a single file and return a Document, or None on failure."""
        try:
            content = self._extract_text(path)
            if not content or not content.strip():
                logger.debug(f"[FilesystemReader] Skipping empty file: {path}")
                return None

            return Document(
                source=str(path.relative_to(self.base_dir)),
                content=content.strip(),
                metadata={
                    "expert_id": expert_id,
                    "source_type": source_type,      # "knowledge_hub" or "master_cases"
                    "filename": path.name,
                    "extension": path.suffix.lower(),
                    "size_bytes": path.stat().st_size,
                },
            )
        except Exception as e:
            logger.error(f"[FilesystemReader] Failed to read {path}: {e}")
            return None

    def _extract_text(self, path: Path) -> str:
        """Extract plain text from a file based on its extension."""
        ext = path.suffix.lower()

        if ext in {".txt", ".md"}:
            return path.read_text(encoding="utf-8", errors="ignore")

        elif ext == ".json":
            import json
            data = json.loads(path.read_text(encoding="utf-8"))
            # Flatten JSON to readable text
            return json.dumps(data, indent=2)

        elif ext == ".pdf":
            return self._extract_pdf(path)

        elif ext == ".docx":
            return self._extract_docx(path)

        return ""

    def _extract_pdf(self, path: Path) -> str:
        """Extract text from a PDF file."""
        try:
            import pypdf
            reader = pypdf.PdfReader(str(path))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except ImportError:
            logger.warning("[FilesystemReader] pypdf not installed. "
                           "Install with: pip install pypdf")
            return ""
        except Exception as e:
            logger.error(f"[FilesystemReader] PDF extraction failed for {path}: {e}")
            return ""

    def _extract_docx(self, path: Path) -> str:
        """Extract text from a DOCX file."""
        try:
            import docx
            doc = docx.Document(str(path))
            return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
        except ImportError:
            logger.warning("[FilesystemReader] python-docx not installed. "
                           "Install with: pip install python-docx")
            return ""
        except Exception as e:
            logger.error(f"[FilesystemReader] DOCX extraction failed for {path}: {e}")
            return ""
