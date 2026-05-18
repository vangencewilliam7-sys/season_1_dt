"""
Document Processor Service
===========================
Extracts plain text from uploaded clinical documents.
Supports: PDF (via PyMuPDF/fitz), TXT, MD
"""

import os
import io
import docx
from typing import Optional


class DocumentProcessor:
    """
    Extracts readable text from clinical documents.
    Uses PyMuPDF for high-fidelity PDF extraction.
    Falls back gracefully for unsupported types.
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}

    @classmethod
    def extract(cls, file_bytes: bytes, filename: str) -> dict:
        """
        Main entry point. Returns a dict with:
          - text: str          — full extracted text
          - filename: str      — original filename
          - size_bytes: int    — size of the uploaded file
          - page_count: int    — number of pages (PDFs only)
          - format: str        — detected format
          - error: str|None    — error message if extraction failed
        """
        filename_lower = filename.lower()
        ext = os.path.splitext(filename_lower)[1]
        size_bytes = len(file_bytes)

        result = {
            "filename": filename,
            "size_bytes": size_bytes,
            "format": ext.lstrip(".").upper() or "UNKNOWN",
            "page_count": 1,
            "text": "",
            "error": None,
        }

        try:
            if ext == ".pdf":
                result.update(cls._extract_pdf(file_bytes))
            elif ext == ".docx":
                result.update(cls._extract_docx(file_bytes))
            elif ext in (".txt", ".md"):
                result.update(cls._extract_text(file_bytes))
            else:
                result["error"] = (
                    f"Unsupported file type '{ext}'. "
                    f"Supported types: {', '.join(cls.SUPPORTED_EXTENSIONS)}"
                )
        except Exception as e:
            result["error"] = f"Extraction failed: {str(e)}"

        return result

    # ------------------------------------------------------------------
    # Private helpers
    # ------------------------------------------------------------------

    @classmethod
    def _extract_docx(cls, file_bytes: bytes) -> dict:
        """Extract text from DOCX using python-docx."""
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            
            text = "\n".join(full_text).strip()
            if not text:
                text = "[DOCX contained no extractable text]"
                
            return {"text": text, "page_count": 1, "format": "DOCX"}
        except Exception as e:
            raise ValueError(f"DOCX extraction failed: {str(e)}")

    @classmethod
    def _extract_pdf(cls, file_bytes: bytes) -> dict:
        """Extract text from PDF using PyMuPDF (fitz)."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError(
                "PyMuPDF is not installed. Run: pip install pymupdf"
            )

        pages_text = []
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            page_count = len(doc)
            for page in doc:
                pages_text.append(page.get_text("text"))

        full_text = "\n\n".join(pages_text).strip()

        if not full_text:
            full_text = "[PDF contained no extractable text — may be image-based or scanned]"

        return {"text": full_text, "page_count": page_count, "format": "PDF"}

    @classmethod
    def _extract_text(cls, file_bytes: bytes) -> dict:
        """Decode plain text / markdown files."""
        for encoding in ("utf-8", "latin-1", "cp1252"):
            try:
                text = file_bytes.decode(encoding).strip()
                return {"text": text, "page_count": 1, "format": "TXT"}
            except UnicodeDecodeError:
                continue
        raise ValueError("Could not decode file — unknown encoding.")

    @classmethod
    def truncate(cls, text: str, max_chars: int = 8000) -> str:
        """
        Truncate extracted text to keep LLM context window manageable.
        Adds a notice if content was trimmed.
        """
        if len(text) <= max_chars:
            return text
        return (
            text[:max_chars]
            + f"\n\n[... Document truncated to {max_chars} characters for processing ...]"
        )
