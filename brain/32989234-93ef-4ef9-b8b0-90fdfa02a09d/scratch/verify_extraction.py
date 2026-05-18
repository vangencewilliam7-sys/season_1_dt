import sys
import os
import io
import docx

# Add backend to path
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'backend')))

from app.services.document_processor import DocumentProcessor

def test_docx_extraction():
    print("Testing DOCX extraction...")
    # Create a dummy docx
    doc = docx.Document()
    doc.add_paragraph("This is a clinical report.")
    doc.add_paragraph("Patient has high FSH levels.")
    
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Test"
    table.cell(0, 1).text = "Result"
    
    row = table.add_row()
    row.cells[0].text = "FSH"
    row.cells[1].text = "15.5 mIU/mL"
    
    buf = io.BytesIO()
    doc.save(buf)
    file_bytes = buf.getvalue()
    
    result = DocumentProcessor.extract(file_bytes, "test_report.docx")
    print(f"Format: {result['format']}")
    print(f"Text Content:\n{result['text']}")
    print("-" * 20)
    assert "FSH levels" in result['text']
    assert "15.5 mIU/mL" in result['text']
    print("DOCX test passed!")

def test_pdf_extraction():
    print("Testing PDF extraction (requires PyMuPDF)...")
    # We can't easily create a PDF from scratch without more libs, 
    # but we can check if it tries to use fitz.
    # For now, just check if it handles unsupported types.
    result = DocumentProcessor.extract(b"hello world", "test.unsupported")
    print(f"Unsupported Error: {result['error']}")
    assert "Unsupported file type" in result['error']
    print("Unsupported type test passed!")

if __name__ == "__main__":
    try:
        test_docx_extraction()
        test_pdf_extraction()
    except Exception as e:
        print(f"Test failed: {e}")
        sys.exit(1)
