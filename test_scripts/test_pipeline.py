import sys
import os
import asyncio

sys.path.append(os.path.join(os.path.dirname(__file__), "knnowledge_Hub", "backend"))

from app.graph.pipeline import create_pipeline
from app.models.state import GraphState

def run():
    print("Initializing pipeline...")
    pipeline = create_pipeline()
    
    # We will pass a dummy source_path that actually exists.
    test_file = os.path.join(os.path.dirname(__file__), "knnowledge_Hub", "backend", "requirements.txt")
    
    # Wait, we need to test with an actual PDF or DOCX file to hit the parser.
    # Let's create a small dummy docx.
    from docx import Document
    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a test document.')
    dummy_docx = os.path.join(os.path.dirname(__file__), "dummy_test.docx")
    doc.save(dummy_docx)
    
    initial_state = GraphState(document_id="test-123", source_path=dummy_docx)
    
    print("Invoking pipeline...")
    try:
        result = pipeline.invoke(initial_state, config={"configurable": {"thread_id": "test-123"}})
        print("Pipeline finished successfully!")
        print(f"Status: {result.get('pipeline_status', 'Unknown')}")
    except Exception as e:
        print(f"Pipeline crashed! {e}")

if __name__ == "__main__":
    run()
